"""Career Pro V2 helpers: upload extraction and truthful resume optimization."""
from __future__ import annotations
import io, json, os, re
import requests
from career_engine import analyze_resume
MAX_UPLOAD_BYTES=8*1024*1024

def extract_resume_upload(f):
    name=(f.filename or '').lower(); raw=f.read(MAX_UPLOAD_BYTES+1)
    if not raw: raise ValueError('The uploaded resume is empty.')
    if len(raw)>MAX_UPLOAD_BYTES: raise ValueError('Resume file is too large. Please keep it under 8 MB.')
    if name.endswith('.pdf'):
        from pypdf import PdfReader
        text='\n\n'.join((p.extract_text() or '') for p in PdfReader(io.BytesIO(raw)).pages).strip()
    elif name.endswith('.docx'):
        from docx import Document
        d=Document(io.BytesIO(raw)); parts=[p.text.strip() for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                line=' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                if line: parts.append(line)
        text='\n'.join(parts).strip()
    elif name.endswith('.txt'): text=raw.decode('utf-8',errors='replace').strip()
    else: raise ValueError('Please upload a PDF, DOCX, or TXT resume.')
    if len(text)<40: raise ValueError("We couldn't extract enough readable text. If this is a scanned PDF, upload a text-based PDF or DOCX instead.")
    return text

def _fallback(resume,job):
    a=analyze_resume(resume,job); k=a['keyword_analysis']; skills=a.get('skills_detected',[])[:12]
    return {'mode':'structured','current_match':k['match_percentage'],'optimized_match':None,'optimized_summary':'Use the strongest truthful role-relevant experience from your existing summary and experience. Add target-job language only where it accurately describes work you have done.','optimized_bullets':[],'skills_to_highlight':k['matched'][:12] or skills,'keywords_to_review':k['missing'][:12],'optimized_resume':resume,'changes':['Prioritize matched job language in the summary and experience bullets.','Remove unrelated wording where it weakens relevance.','Add missing keywords only when they truthfully describe your experience.'],'integrity_note':'No new employers, dates, education, metrics, skills, certifications, or achievements should be added unless supported by your resume.'}

def optimize_with_gemini(resume,job):
    key=(os.getenv('GEMINI_API') or os.getenv('GEMINI_API_KEY') or '').strip()
    if not key:return _fallback(resume,job)
    model=os.getenv('GEMINI_MODEL','gemini-2.5-flash')
    prompt=f'''You are EggyPDF Career Pro. Optimize this resume for the target job. NEVER invent facts. Preserve employers, titles, dates, education and contact facts. Never invent metrics, certifications, tools, skills, responsibilities or achievements. Rewrite/reorder existing facts and use a job keyword only when the resume provides evidence. Return JSON only with keys optimized_summary (string), optimized_bullets (array of strings), skills_to_highlight (array), keywords_to_review (array), optimized_resume (string), changes (array of strings).\nRESUME:\n{resume}\n\nTARGET JOB:\n{job}'''
    try:
        r=requests.post(f'https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}',json={'contents':[{'parts':[{'text':prompt}]}],'generationConfig':{'responseMimeType':'application/json','temperature':0.2}},timeout=35);r.raise_for_status()
        text=r.json()['candidates'][0]['content']['parts'][0]['text'].strip();text=re.sub(r'^```(?:json)?\s*|\s*```$','',text,flags=re.I|re.S);out=json.loads(text)
        before=analyze_resume(resume,job)['keyword_analysis']['match_percentage'];after=analyze_resume(out.get('optimized_resume') or resume,job)['keyword_analysis']['match_percentage']
        out.update({'mode':'ai','current_match':before,'optimized_match':after,'integrity_note':'Review AI rewrites before use. EggyPDF is instructed not to invent experience, metrics, skills, education, employers, or achievements.'});return out
    except Exception:return _fallback(resume,job)
